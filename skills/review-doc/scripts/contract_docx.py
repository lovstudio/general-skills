#!/usr/bin/env python3
"""Extract, validate, and apply contract comments/revisions to DOCX files.

The script uses stable block anchors for body paragraphs, table cells, headers,
and footers. It deliberately rejects unsafe cross-run revisions instead of
silently damaging formatting.
"""

from __future__ import annotations

import argparse
import copy
import json
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document
    from docx.document import Document as DocumentObject
    from docx.oxml import parse_xml
    from docx.opc.packuri import PackURI
    from docx.opc.part import XmlPart
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    from lxml import etree
except ImportError as exc:  # pragma: no cover - dependency failure path
    raise SystemExit(
        "缺少依赖。请在隔离环境安装：python3 -m pip install 'python-docx>=1.0' 'lxml>=4.9'"
    ) from exc


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
NSMAP = {"w": W_NS}
VALID_SEVERITIES = {"致命", "高", "中", "低", "提示"}
VALID_PRIORITIES = {"必须坚持", "优先争取", "可交换", "可接受"}


def qn(tag: str) -> str:
    prefix, local = tag.split(":", 1)
    return f"{{{NSMAP[prefix]}}}{local}"


@dataclass(frozen=True)
class Block:
    anchor: str
    location: str
    paragraph: Paragraph

    @property
    def text(self) -> str:
        return visible_text(self.paragraph._element)


def visible_text(paragraph_element: etree._Element) -> str:
    """Return visible text, including insertions and excluding deletions."""
    chunks: list[str] = []
    for element in paragraph_element.iter():
        if element.tag == qn("w:t") and not _has_ancestor(element, qn("w:del")):
            chunks.append(element.text or "")
        elif element.tag == qn("w:tab") and not _has_ancestor(element, qn("w:del")):
            chunks.append("\t")
        elif element.tag in {qn("w:br"), qn("w:cr")} and not _has_ancestor(
            element, qn("w:del")
        ):
            chunks.append("\n")
    return "".join(chunks).strip()


def _has_ancestor(element: etree._Element, tag: str) -> bool:
    parent = element.getparent()
    while parent is not None:
        if parent.tag == tag:
            return True
        parent = parent.getparent()
    return False


def _walk_container(
    element: etree._Element,
    parent: DocumentObject | _Cell | Any,
    prefix: str,
) -> Iterable[Block]:
    paragraph_index = 0
    table_index = 0
    for child in element.iterchildren():
        if child.tag == qn("w:p"):
            anchor = f"{prefix}.p{paragraph_index}"
            yield Block(anchor, anchor, Paragraph(child, parent))
            paragraph_index += 1
            continue

        if child.tag != qn("w:tbl"):
            continue

        table = Table(child, parent)
        current_table = table_index
        table_index += 1
        seen_cells: set[etree._Element] = set()
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                cell_prefix = (
                    f"{prefix}.tbl{current_table}.r{row_index}.c{cell_index}"
                )
                yield from _walk_container(cell._tc, cell, cell_prefix)


def collect_blocks(document: DocumentObject) -> list[Block]:
    blocks = list(_walk_container(document.element.body, document, "body"))

    seen_parts: set[str] = set()
    variants = (
        ("header", "header"),
        ("first-header", "first_page_header"),
        ("even-header", "even_page_header"),
        ("footer", "footer"),
        ("first-footer", "first_page_footer"),
        ("even-footer", "even_page_footer"),
    )
    for section_index, section in enumerate(document.sections):
        for label, attribute in variants:
            container = getattr(section, attribute)
            part_name = str(container.part.partname)
            if part_name in seen_parts:
                continue
            seen_parts.add(part_name)
            prefix = f"section{section_index}.{label}"
            blocks.extend(_walk_container(container._element, container, prefix))
    return blocks


def block_map(document: DocumentObject) -> dict[str, Block]:
    return {block.anchor: block for block in collect_blocks(document)}


def load_json(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"无法读取 JSON：{path}：{exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("批注数据的根节点必须是 JSON 对象")
    return data


def _visible_runs(paragraph: Paragraph) -> list[tuple[etree._Element, str]]:
    result: list[tuple[etree._Element, str]] = []
    for run in paragraph._element.iter(qn("w:r")):
        if _has_ancestor(run, qn("w:del")):
            continue
        text = "".join((node.text or "") for node in run.iter(qn("w:t")))
        if text:
            result.append((run, text))
    return result


def _revision_run(paragraph: Paragraph, old: str) -> etree._Element | None:
    matches = [run for run, text in _visible_runs(paragraph) if old in text]
    if len(matches) != 1:
        return None
    run = matches[0]
    if _has_ancestor(run, qn("w:ins")):
        return None
    allowed = {qn("w:rPr"), qn("w:t")}
    if any(child.tag not in allowed for child in run):
        return None
    return run


def validate_data(
    document: DocumentObject, data: dict[str, Any]
) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    blocks = block_map(document)

    comments = data.get("comments", [])
    revisions = data.get("revisions", [])
    if not isinstance(comments, list):
        errors.append("comments 必须是数组")
        comments = []
    if not isinstance(revisions, list):
        errors.append("revisions 必须是数组")
        revisions = []
    if not comments and not revisions:
        warnings.append("没有批注或修订")

    for index, comment in enumerate(comments, start=1):
        prefix = f"comments[{index}]"
        if not isinstance(comment, dict):
            errors.append(f"{prefix} 必须是对象")
            continue
        anchor = comment.get("anchor")
        if anchor not in blocks:
            errors.append(f"{prefix}.anchor 不存在：{anchor!r}")
            continue
        text = comment.get("text")
        if not isinstance(text, str) or not text.strip():
            errors.append(f"{prefix}.text 不能为空")
        severity = comment.get("severity")
        if severity not in VALID_SEVERITIES:
            errors.append(f"{prefix}.severity 必须是 {sorted(VALID_SEVERITIES)} 之一")
        priority = comment.get("priority")
        if priority not in VALID_PRIORITIES:
            errors.append(f"{prefix}.priority 必须是 {sorted(VALID_PRIORITIES)} 之一")
        if not isinstance(comment.get("category"), str) or not comment["category"].strip():
            errors.append(f"{prefix}.category 不能为空")

        quote = comment.get("quote")
        if quote is not None:
            if not isinstance(quote, str) or not quote:
                errors.append(f"{prefix}.quote 必须是非空字符串")
            else:
                count = blocks[anchor].text.count(quote)
                occurrence = comment.get("occurrence", 1)
                if not isinstance(occurrence, int) or occurrence < 1:
                    errors.append(f"{prefix}.occurrence 必须是从 1 开始的整数")
                elif count < occurrence:
                    errors.append(
                        f"{prefix}.quote 在 {anchor} 中仅出现 {count} 次，无法定位第 {occurrence} 次"
                    )
                elif count > 1 and "occurrence" not in comment:
                    warnings.append(
                        f"{prefix}.quote 在 {anchor} 中出现 {count} 次，将定位第一次；建议填写 occurrence"
                    )

    for index, revision in enumerate(revisions, start=1):
        prefix = f"revisions[{index}]"
        if not isinstance(revision, dict):
            errors.append(f"{prefix} 必须是对象")
            continue
        anchor = revision.get("anchor")
        if anchor not in blocks:
            errors.append(f"{prefix}.anchor 不存在：{anchor!r}")
            continue
        old = revision.get("old")
        new = revision.get("new")
        if not isinstance(old, str) or not old:
            errors.append(f"{prefix}.old 不能为空")
            continue
        if not isinstance(new, str):
            errors.append(f"{prefix}.new 必须是字符串")
            continue
        if old == new:
            errors.append(f"{prefix}.new 不得与 old 相同")
        if "\n" in new or "\r" in new:
            errors.append(f"{prefix}.new 不得包含换行；请按段落拆分修订")
        count = blocks[anchor].text.count(old)
        if count != 1:
            errors.append(f"{prefix}.old 在 {anchor} 中应唯一出现，实际出现 {count} 次")
        elif _revision_run(blocks[anchor].paragraph, old) is None:
            errors.append(
                f"{prefix}.old 跨越多个 Word 文本运行、位于既有修订中或含复杂格式；请缩短替换范围或改用批注"
            )
        if not isinstance(revision.get("reason"), str) or not revision["reason"].strip():
            errors.append(f"{prefix}.reason 不能为空")

    return errors, warnings


def _ensure_comments_part(document: DocumentObject) -> Any:
    for relationship in document.part.rels.values():
        if relationship.reltype == COMMENTS_REL:
            return relationship.target_part

    try:
        from docx.parts.comments import CommentsPart

        comments_part = CommentsPart.default(document.part.package)
    except (ImportError, AttributeError):  # pragma: no cover - old python-docx
        root = parse_xml(
            (
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main"/>'
            ).encode("utf-8")
        )
        comments_part = XmlPart(
            PackURI("/word/comments.xml"),
            COMMENTS_CONTENT_TYPE,
            root,
            document.part.package,
        )
    document.part.relate_to(comments_part, COMMENTS_REL)
    return comments_part


def _next_comment_id(document: DocumentObject, comments_part: Any) -> int:
    ids: list[int] = []
    for element in comments_part._element.iter(qn("w:comment")):
        value = element.get(qn("w:id"))
        if value is not None and value.isdigit():
            ids.append(int(value))
    for block in collect_blocks(document):
        for tag in ("w:commentRangeStart", "w:commentRangeEnd"):
            for element in block.paragraph._element.iter(qn(tag)):
                value = element.get(qn("w:id"))
                if value is not None and value.isdigit():
                    ids.append(int(value))
    return max(ids, default=-1) + 1


def _append_comment(
    comments_part: Any, comment_id: int, author: str, text: str, date: str
) -> None:
    comment = etree.SubElement(comments_part._element, qn("w:comment"))
    comment.set(qn("w:id"), str(comment_id))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:date"), date)
    paragraph = etree.SubElement(comment, qn("w:p"))
    run = etree.SubElement(paragraph, qn("w:r"))
    text_element = etree.SubElement(run, qn("w:t"))
    text_element.text = text
    text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _comment_markers(comment_id: int) -> tuple[etree._Element, etree._Element, etree._Element]:
    start = etree.Element(qn("w:commentRangeStart"))
    start.set(qn("w:id"), str(comment_id))
    end = etree.Element(qn("w:commentRangeEnd"))
    end.set(qn("w:id"), str(comment_id))
    reference_run = etree.Element(qn("w:r"))
    run_properties = etree.SubElement(reference_run, qn("w:rPr"))
    style = etree.SubElement(run_properties, qn("w:rStyle"))
    style.set(qn("w:val"), "CommentReference")
    reference = etree.SubElement(reference_run, qn("w:commentReference"))
    reference.set(qn("w:id"), str(comment_id))
    return start, end, reference_run


def _anchor_comment(
    paragraph: Paragraph, comment_id: int, quote: str | None, occurrence: int
) -> None:
    start, end, reference = _comment_markers(comment_id)
    runs = _visible_runs(paragraph)

    if quote and runs:
        combined = "".join(text for _, text in runs)
        start_offset = -1
        cursor = 0
        for _ in range(occurrence):
            start_offset = combined.find(quote, cursor)
            cursor = start_offset + len(quote)
        end_offset = start_offset + len(quote)
        position = 0
        start_run: etree._Element | None = None
        end_run: etree._Element | None = None
        for run, run_text in runs:
            run_end = position + len(run_text)
            if start_run is None and position <= start_offset < run_end:
                start_run = run
            if position < end_offset <= run_end:
                end_run = run
                break
            position = run_end
        if (
            start_run is not None
            and end_run is not None
            and start_run.getparent() is end_run.getparent()
        ):
            start_run.addprevious(start)
            end_run.addnext(end)
            end.addnext(reference)
            return

    paragraph_element = paragraph._element
    insert_at = 1 if len(paragraph_element) and paragraph_element[0].tag == qn("w:pPr") else 0
    paragraph_element.insert(insert_at, start)
    paragraph_element.append(end)
    paragraph_element.append(reference)


def _regular_run(text: str, run_properties: etree._Element | None) -> etree._Element:
    run = etree.Element(qn("w:r"))
    if run_properties is not None:
        run.append(copy.deepcopy(run_properties))
    text_element = etree.SubElement(run, qn("w:t"))
    text_element.text = text
    text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return run


def _tracked_element(
    tag: str,
    text_tag: str,
    text: str,
    author: str,
    date: str,
    revision_id: int,
    run_properties: etree._Element | None,
) -> etree._Element:
    element = etree.Element(qn(tag))
    element.set(qn("w:id"), str(revision_id))
    element.set(qn("w:author"), author)
    element.set(qn("w:date"), date)
    run = etree.SubElement(element, qn("w:r"))
    if run_properties is not None:
        run.append(copy.deepcopy(run_properties))
    text_element = etree.SubElement(run, qn(text_tag))
    text_element.text = text
    text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return element


def _apply_revision(
    paragraph: Paragraph,
    old: str,
    new: str,
    author: str,
    date: str,
    revision_id: int,
) -> None:
    run = _revision_run(paragraph, old)
    if run is None:  # validation should make this unreachable
        raise ValueError(f"无法安全修订：{old!r}")
    run_text = "".join((node.text or "") for node in run.iter(qn("w:t")))
    before, after = run_text.split(old, 1)
    run_properties = run.find(qn("w:rPr"))
    parent = run.getparent()
    position = parent.index(run)
    parent.remove(run)

    elements: list[etree._Element] = []
    if before:
        elements.append(_regular_run(before, run_properties))
    elements.append(
        _tracked_element(
            "w:del", "w:delText", old, author, date, revision_id, run_properties
        )
    )
    elements.append(
        _tracked_element(
            "w:ins", "w:t", new, author, date, revision_id + 1, run_properties
        )
    )
    if after:
        elements.append(_regular_run(after, run_properties))
    for offset, element in enumerate(elements):
        parent.insert(position + offset, element)


def _next_revision_id(document: DocumentObject) -> int:
    ids: list[int] = []
    for block in collect_blocks(document):
        for tag in ("w:ins", "w:del"):
            for element in block.paragraph._element.iter(qn(tag)):
                value = element.get(qn("w:id"))
                if value is not None and value.isdigit():
                    ids.append(int(value))
    return max(ids, default=999) + 1


def apply_annotations(
    document: DocumentObject, data: dict[str, Any], output_path: Path
) -> dict[str, Any]:
    blocks = block_map(document)
    metadata = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    default_author = str(metadata.get("reviewer") or "合同审阅专家")
    now = datetime.now().astimezone().isoformat(timespec="seconds")

    comments = data.get("comments", [])
    if comments:
        comments_part = _ensure_comments_part(document)
        comment_id = _next_comment_id(document, comments_part)
        for comment in comments:
            author = str(comment.get("author") or default_author)
            _append_comment(comments_part, comment_id, author, comment["text"], now)
            _anchor_comment(
                blocks[comment["anchor"]].paragraph,
                comment_id,
                comment.get("quote"),
                int(comment.get("occurrence", 1)),
            )
            comment_id += 1

    revision_id = _next_revision_id(document)
    revisions = data.get("revisions", [])
    for revision in revisions:
        author = str(revision.get("author") or default_author)
        _apply_revision(
            blocks[revision["anchor"]].paragraph,
            revision["old"],
            revision["new"],
            author,
            now,
            revision_id,
        )
        revision_id += 2

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return {
        "output": str(output_path),
        "comments": len(comments),
        "revisions": len(revisions),
    }


def extract_command(input_path: Path, output_path: Path | None) -> None:
    document = Document(input_path)
    blocks = collect_blocks(document)
    payload = {
        "source": str(input_path),
        "blocks": [
            {
                "anchor": block.anchor,
                "location": block.location,
                "text": block.text,
            }
            for block in blocks
            if block.text
        ],
        "limitations": [
            "不提取文本框、脚注、尾注、嵌入对象或既有批注正文",
            "修订删除内容不计入可见文本，修订插入内容计入可见文本",
        ],
    }
    serialized = json.dumps(payload, ensure_ascii=False, indent=2) + "\n"
    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(serialized, encoding="utf-8")
        print(json.dumps({"output": str(output_path), "blocks": len(payload["blocks"])}, ensure_ascii=False))
    else:
        sys.stdout.write(serialized)


def validate_command(input_path: Path, annotations_path: Path) -> int:
    document = Document(input_path)
    data = load_json(annotations_path)
    errors, warnings = validate_data(document, data)
    print(
        json.dumps(
            {"valid": not errors, "errors": errors, "warnings": warnings},
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if not errors else 2


def annotate_command(
    input_path: Path, annotations_path: Path, output_path: Path
) -> int:
    if input_path.resolve() == output_path.resolve():
        print("错误：输出路径不得覆盖原文件", file=sys.stderr)
        return 2
    document = Document(input_path)
    data = load_json(annotations_path)
    errors, warnings = validate_data(document, data)
    for warning in warnings:
        print(f"警告：{warning}", file=sys.stderr)
    if errors:
        print(json.dumps({"valid": False, "errors": errors}, ensure_ascii=False, indent=2), file=sys.stderr)
        return 2
    result = apply_annotations(document, data, output_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="提取合同 DOCX、校验审阅数据并写入批注与修订"
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract_parser = subparsers.add_parser("extract", help="提取可审阅文本块")
    extract_parser.add_argument("--input", type=Path, required=True)
    extract_parser.add_argument("--output", type=Path)

    validate_parser = subparsers.add_parser("validate", help="校验批注和修订 JSON")
    validate_parser.add_argument("--input", type=Path, required=True)
    validate_parser.add_argument("--annotations", type=Path, required=True)

    annotate_parser = subparsers.add_parser("annotate", help="写入批注和修订")
    annotate_parser.add_argument("--input", type=Path, required=True)
    annotate_parser.add_argument("--annotations", type=Path, required=True)
    annotate_parser.add_argument("--output", type=Path, required=True)

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        if args.command == "extract":
            extract_command(args.input, args.output)
            return 0
        if args.command == "validate":
            return validate_command(args.input, args.annotations)
        if args.command == "annotate":
            return annotate_command(args.input, args.annotations, args.output)
    except (OSError, ValueError, KeyError) as exc:
        print(f"错误：{exc}", file=sys.stderr)
        return 2
    parser.error("未知命令")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
