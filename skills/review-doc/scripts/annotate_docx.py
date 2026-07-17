#!/usr/bin/env python3
"""review-doc 0.x compatibility entrypoint.

New integrations should use ``contract_docx.py`` and stable ``anchor`` fields.
This wrapper keeps the former paragraph-index JSON accepted for existing users.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document

import contract_docx


def extract_legacy(input_path: Path) -> None:
    document = Document(input_path)
    result = []
    for index, paragraph in enumerate(document.paragraphs):
        text = contract_docx.visible_text(paragraph._element)
        if text:
            result.append({"index": index, "text": text})
    json.dump(result, sys.stdout, ensure_ascii=False, indent=2)
    sys.stdout.write("\n")


def _legacy_quote(comment: dict[str, Any], paragraph_text: str) -> str | None:
    start = comment.get("start")
    end = comment.get("end")
    if not isinstance(start, int) or not isinstance(end, int):
        return None
    if start < 0 or end <= start or end > len(paragraph_text):
        return None
    return paragraph_text[start:end]


def convert_legacy(document: Any, data: dict[str, Any]) -> dict[str, Any]:
    comments = data.get("comments", [])
    revisions = data.get("revisions", [])
    if any(isinstance(item, dict) and "anchor" in item for item in comments + revisions):
        return data

    reviewer = "合同审阅专家"
    for item in comments + revisions:
        if isinstance(item, dict) and item.get("author"):
            reviewer = str(item["author"])
            break

    converted_comments = []
    for item in comments:
        if not isinstance(item, dict) or "paragraph" not in item:
            continue
        paragraph_index = int(item["paragraph"])
        paragraph_text = (
            contract_docx.visible_text(document.paragraphs[paragraph_index]._element)
            if 0 <= paragraph_index < len(document.paragraphs)
            else ""
        )
        converted = {
            "anchor": f"body.p{paragraph_index}",
            "severity": item.get("severity", "提示"),
            "category": item.get("category", "兼容导入"),
            "priority": item.get("priority", "可接受"),
            "text": item.get("text", ""),
            "author": item.get("author", reviewer),
        }
        quote = _legacy_quote(item, paragraph_text)
        if quote:
            converted["quote"] = quote
        converted_comments.append(converted)

    converted_revisions = []
    for item in revisions:
        if not isinstance(item, dict) or "paragraph" not in item:
            continue
        converted_revisions.append(
            {
                "anchor": f"body.p{int(item['paragraph'])}",
                "old": item.get("old", ""),
                "new": item.get("new", ""),
                "reason": item.get("reason", "兼容旧版段落索引修订"),
                "author": item.get("author", reviewer),
            }
        )

    return {
        "metadata": {
            "reviewer": reviewer,
            "migration": "review-doc-0.x-paragraph-index",
        },
        "comments": converted_comments,
        "revisions": converted_revisions,
    }


def annotate_legacy(input_path: Path, annotations_path: Path, output_path: Path) -> int:
    document = Document(input_path)
    data = contract_docx.load_json(annotations_path)
    converted = convert_legacy(document, data)
    errors, warnings = contract_docx.validate_data(document, converted)
    for warning in warnings:
        print(f"警告：{warning}", file=sys.stderr)
    if errors:
        print(
            json.dumps({"valid": False, "errors": errors}, ensure_ascii=False, indent=2),
            file=sys.stderr,
        )
        return 2
    if input_path.resolve() == output_path.resolve():
        print("错误：输出路径不得覆盖原文件", file=sys.stderr)
        return 2
    result = contract_docx.apply_annotations(document, converted, output_path)
    result["compatibility_mode"] = True
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="review-doc 0.x 兼容入口")
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract_parser = subparsers.add_parser("extract")
    extract_parser.add_argument("--input", type=Path, required=True)

    annotate_parser = subparsers.add_parser("annotate")
    annotate_parser.add_argument("--input", type=Path, required=True)
    annotate_parser.add_argument("--annotations", type=Path, required=True)
    annotate_parser.add_argument("--output", type=Path, required=True)
    return parser


def main() -> int:
    args = build_parser().parse_args()
    try:
        if args.command == "extract":
            extract_legacy(args.input)
            return 0
        return annotate_legacy(args.input, args.annotations, args.output)
    except (OSError, ValueError, IndexError, KeyError) as exc:
        print(f"错误：{exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
